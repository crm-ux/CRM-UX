# -*- coding: utf-8 -*-
from odoo import api, fields, models
from odoo.tools import html2plaintext
from markupsafe import Markup


class SaleQuotePreviewWizard(models.TransientModel):
    _name = 'sale.quote.preview.wizard'
    _description = 'Editable Quotation Preview Wizard'

    order_id = fields.Many2one('sale.order', string='Quotation', required=True, readonly=True)
    x_gst_included = fields.Boolean(string='Include GST', default=True)
    seller_name = fields.Char(string='Seller / Company Name')
    buyer_name = fields.Char(string='Party Name')
    quote_name = fields.Char(string='Quotation No.')
    quote_date = fields.Date(string='Quotation Date')
    valid_until = fields.Date(string='Valid Upto')
    subject = fields.Char(string='Subject')
    best_offer_for = fields.Char(string='Best Offer For', help='Text after best offer for in the intro paragraph')

    company_logo = fields.Binary(string='Company Logo')
    document_html = fields.Html(string='Editable Quote Content')
    technical_specs_html = fields.Html(string='Technical Specifications')
    quote_image_ids = fields.Many2many(
        'ir.attachment', 'sale_quote_wizard_image_rel',
        'wizard_id', 'attachment_id',
        string='Quote Images'
    )

    @api.model
    def default_get(self, fields_list):
        res = super().default_get(fields_list)
        order = self.env['sale.order'].browse(self.env.context.get('default_order_id'))
        if not order:
            return res

        gst_on = order.x_gst_included

        rows = ''
        for idx, line in enumerate(order.order_line.filtered(lambda l: not l.display_type), 1):
            part_no = line.x_product_code or line.product_id.default_code or ''
            note = (line.x_notes if hasattr(line, 'x_notes') and line.x_notes else '')
            hsn = line.product_id.l10n_in_hsn_code or ''
            make = line.x_make or ''
            unit_price = line.price_unit or 0
            discount_pct = line.discount or 0
            disc_amount = unit_price * discount_pct / 100
            after_discount = unit_price - disc_amount
            qty = line.product_uom_qty or 0
            amount = line.price_subtotal or 0

            # Tax
            tax_str = ''
            for tax in line.tax_ids:
                tname = (tax.name or '').upper()
                tax_str += '%s(%s%%) ' % (tname, int(tax.amount))
            tax_str = tax_str.strip() or '-'

            # Discount
            if discount_pct:
                disc_str = '(%s%%)=%s' % (int(discount_pct), int(disc_amount)) if discount_pct else '-'
            else:
                disc_str = '-'

            # Description
            desc_html = line.x_product_name or line.product_id.name or ''
            if part_no:
                desc_html += '<br/><small style="color:#666;">Part No: %s</small>' % part_no
            if make:
                desc_html += '<br/><small style="color:#555;">Make: %s</small>' % make
            if note:
                desc_html += '<br/><small style="color:#888;font-style:italic;">Note: %s</small>' % note

            # Image
            img_html = ''
            if line.product_id.image_128:
                img_b64 = line.product_id.image_128.decode('utf-8') if isinstance(line.product_id.image_128, bytes) else line.product_id.image_128
                img_html = ""

            tax_cell = '<td style="text-align:center;" contenteditable="true">%s</td>' % tax_str if gst_on else ''

            rows += (
                '<tr>'
                '<td style="text-align:center;">%s</td>'
                '<td>%s%s</td>'
                '<td style="text-align:center;">%s</td>'
                '<td style="text-align:right;" contenteditable="false">%s</td>'
                '<td style="text-align:right;" contenteditable="false">%s</td>'
                '<td style="text-align:right;" contenteditable="false">%s</td>'
                '<td style="text-align:right;" contenteditable="false">%s</td>'
                '<td style="text-align:right;" contenteditable="false">%s</td>'
                '</tr>'
            ) % (idx, desc_html, img_html, hsn, int(unit_price), disc_str, int(after_discount), int(qty), int(amount))

        # Technical Specs Page
        tech_rows = ''
        for idx, line in enumerate(order.order_line.filtered(lambda l: not l.display_type), 1):
            if line.x_technical_specs:
                pname = line.x_product_name or line.product_id.name or ''
                tech_rows += (
                    '<h4>%s. %s</h4>'
                    '<p>%s</p>'
                ) % (idx, pname, line.x_technical_specs.replace('\n', '<br/>'))

        # Images Page
        img_rows = ''
        for idx, line in enumerate(order.order_line.filtered(lambda l: not l.display_type), 1):
            img_src = ''
            if line.x_product_image:
                img_b64 = line.x_product_image.decode('utf-8') if isinstance(line.x_product_image, bytes) else line.x_product_image
                img_src = '<img src="data:image/png;base64,%s" style="max-width:300px;max-height:300px;"/>' % img_b64
            elif line.product_id.image_1920:
                img_b64 = line.product_id.image_1920.decode('utf-8') if isinstance(line.product_id.image_1920, bytes) else line.product_id.image_1920
                img_src = '<img src="data:image/png;base64,%s" style="max-width:300px;max-height:300px;"/>' % img_b64
            if img_src:
                pname = line.x_product_name or line.product_id.name or ''
                img_rows += '<div style="margin-bottom:20px;"><h4>%s</h4>%s</div>' % (pname, img_src)

        # Get product categories from linked opportunity only
        product_cats = self.env.context.get('best_offer_for', '')
        if not product_cats:
            if order.opportunity_id and order.opportunity_id.x_product_category_ids:
                product_cats = ', '.join(order.opportunity_id.x_product_category_ids.mapped('name'))

        tax_th = '<th style="text-align:center;">Tax</th>' if gst_on else ''
        tax_amount_row = '<p style="text-align:right;"><b>Tax:</b> %s</p>' % int(order.amount_tax) if gst_on else ''

        # Logo
        logo_html = ''
        if order.company_id.logo_web:
            logo_b64 = order.company_id.logo_web.decode('utf-8') if isinstance(order.company_id.logo_web, bytes) else order.company_id.logo_web
            logo_html = '<img src="data:image/png;base64,%s" style="max-height:80px;"/>' % logo_b64

        intro_text = self.env['ir.config_parameter'].sudo().get_param('sale.quote.intro.template', 'With reference to your discussion with the undersigned as regards your subject requirement, we are pleased to quote our best offer for')
        html = (
            '<div style="text-align:right;">%s</div>'
            '<table style="width:100%%;border-collapse:collapse;margin-bottom:8px;background:#eaf0fb;">'
            '<tr>'
            '<td style="padding:6px 12px;font-weight:bold;"><b>Quotation No:</b> %s</td>'
            '<td style="padding:6px 12px;font-weight:bold;text-align:right;"><b>Date:</b> %s</td>'
            '</tr>'
            '</table>'
            '<p><b>To,</b> %s</p>'
            '<p>%s</p>'
            '<p>%s</p>'
            '<p>%s</p>'
            '<p>%s</p>'
            '<br/>'
            '<p><b>Subject:</b> Quotation for Products / Services</p>'
            '<br/>'
            '<p>Dear Sir,</p>'
            '<p>%s</p>'
            '<br/>'
            '<table border="1" cellpadding="6" cellspacing="0" style="width:100%%;border-collapse:collapse;font-size:12px;" contenteditable="false">'
            '<thead><tr style="background:#f0f0f0;">'
            '<th style="text-align:center;width:35px;">SR No.</th>'
            '<th style="text-align:left;">Item Description</th>'
            '<th style="text-align:center;">HSN</th>'
            '<th style="text-align:right;">Unit Price</th>'
            '<th style="text-align:right;">Discount</th>'
            '<th style="text-align:right;">After Discount</th>'
            '<th style="text-align:right;">Qty</th>'
            '<th style="text-align:right;">Amount</th>'
            '</tr></thead>'
            '<tbody>%s</tbody>'
            '</table>'
            '<br/>'
            '<p style="text-align:right;"><b>Gross Total Amount:</b> %s</p>'
            '%s'
            '<p style="text-align:right;font-size:14px;"><b>Total:</b> %s</p>'
            '<h4>Terms &amp; Conditions</h4>'
            '<p>%s</p>'
        ) % (
            logo_html,
            order.name or '',
            order.date_order.date() if order.date_order else fields.Date.today(),
            order.partner_id.name or '',
            order.partner_id.city or '',
            order.partner_id.email or '',
            order.partner_id.phone or '',
            order.partner_id.mobile if hasattr(order.partner_id, 'mobile') else '',
            intro_text,
            rows,
            int(order.amount_untaxed),
            tax_amount_row,
            int(order.amount_total),
            order.note or '',
        )

        # Build tech specs section
        tech_section = ''
        if tech_rows:
            tech_section = (
                '<div style="page-break-before:always;">'
                '<div style="text-align:right;">%s</div>'
                '<h2>Technical Specifications</h2>%s'
                '</div>'
            ) % (logo_html, tech_rows)

        # Build images section
        img_section = ''
        if img_rows:
            img_section = (
                '<div style="page-break-before:always;">'
                '<div style="text-align:right;">%s</div>'
                '<h2>Product Images</h2>%s'
                '</div>'
            ) % (logo_html, img_rows)

        res.update({
            'order_id': order.id,
            'x_gst_included': gst_on,
            'seller_name': order.company_id.name or '',
            'buyer_name': order.partner_id.name or '',
            'quote_name': order.name or '',
            'quote_date': order.date_order.date() if order.date_order else fields.Date.today(),
            'valid_until': order.validity_date,
            'subject': 'Quotation for Products / Services',
            'best_offer_for': getattr(order, 'x_draft_best_offer', None) or product_cats,
            'company_logo': order.company_id.logo_web,
            'document_html': Markup(html),
            'technical_specs_html': getattr(order, 'x_draft_tech_specs', None) or False,
            'quote_image_ids': [(6, 0, getattr(order, 'x_draft_image_ids', self.env['ir.attachment']).ids)],
        })
        return res

    def action_save_draft(self):
        self.ensure_one()
        order = self.order_id
        if not order:
            return
        vals = {
            'x_draft_tech_specs': self.technical_specs_html,
            'x_draft_best_offer': self.best_offer_for,
        }
        if hasattr(order, 'x_draft_image_ids'):
            vals['x_draft_image_ids'] = [(6, 0, self.quote_image_ids.ids)]
        order.sudo().write(vals)
        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'title': 'Draft Saved!',
                'message': 'Your content has been saved. Reopen Preview to continue.',
                'type': 'success',
                'sticky': False,
            }
        }

    def action_add_technical_specs(self):
        self.ensure_one()
        return {
            'type': 'ir.actions.act_window',
            'name': 'Technical Specifications',
            'res_model': 'sale.quote.preview.wizard',
            'res_id': self.id,
            'view_mode': 'form',
            'view_id': self.env.ref('custom_crm_extended.view_quote_tech_specs_form').id,
            'target': 'new',
        }

    def action_save_and_reopen(self):
        self.ensure_one()
        self._rebuild_document_html()
        # Reopen the main preview wizard keeping state
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'sale.quote.preview.wizard',
            'res_id': self.id,
            'view_mode': 'form',
            'target': 'new',
            'flags': {'mode': 'edit'},
        }


    def _style_html_tables(self, html_content):
        """Add proper border and styling to all tables in HTML content"""
        from markupsafe import Markup
        import re
        html_str = str(html_content)
        # Style table tags
        html_str = re.sub(
            r'<table(?![^>]*style)[^>]*>',
            '<table style="width:100%;border-collapse:collapse;font-size:12px;margin-bottom:10px;" border="1" cellpadding="6" cellspacing="0">',
            html_str
        )
        # Style th tags
        html_str = re.sub(
            r'<th(?![^>]*style)[^>]*>',
            '<th style="background:#f0f0f0;border:1px solid #ccc;padding:6px;text-align:left;font-weight:bold;">',
            html_str
        )
        # Style td tags
        html_str = re.sub(
            r'<td(?![^>]*style)[^>]*>',
            '<td style="border:1px solid #ccc;padding:6px;text-align:left;">',
            html_str
        )
        return Markup(html_str)

    def _rebuild_document_html(self):
        from markupsafe import Markup
        order = self.order_id
        if not order:
            return

        # Get base HTML by regenerating
        defaults = self.with_context(default_order_id=order.id, best_offer_for=self.best_offer_for or '').default_get(['document_html'])
        base_html = defaults.get('document_html', '')
        # Replace placeholder with actual best_offer_for if needed
        if self.best_offer_for:
            import re
            from markupsafe import Markup
            base_html_str = str(base_html)
            base_html_str = re.sub(
                r'our best offer for [^<.]*\.',
                'our best offer for %s.' % self.best_offer_for,
                base_html_str
            )
            base_html = Markup(base_html_str)

        # Append tech specs from wizard field
        tech_html = ''
        if self.technical_specs_html:
            logo_html = ''
            if order.company_id.logo_web:
                logo_b64 = order.company_id.logo_web.decode('utf-8') if isinstance(order.company_id.logo_web, bytes) else order.company_id.logo_web
                logo_html = '<img src="data:image/png;base64,%s" style="max-height:80px;"/>' % logo_b64
            styled_specs = self._style_html_tables(self.technical_specs_html)
            tech_html = (
                '<div style="page-break-before:always;">'
                '<div style="text-align:right;">%s</div>'
                '<h2 style="margin-bottom:12px;">Technical Specifications</h2>'
                '<style>'
                'table{width:100%%;border-collapse:collapse;font-size:12px;margin-bottom:10px;}'
                'th{background:#f0f0f0;border:1px solid #999;padding:8px;text-align:left;font-weight:bold;}'
                'td{border:1px solid #999;padding:8px;text-align:left;}'
                'tr:nth-child(even){background:#f9f9f9;}'
                '</style>'
                '<div style="font-size:12px;">%s</div>'
                '</div>'
            ) % (logo_html, styled_specs)

        # Append images from wizard field
        img_html = ''
        if self.quote_image_ids:
            logo_html = ''
            if order.company_id.logo_web:
                logo_b64 = order.company_id.logo_web.decode('utf-8') if isinstance(order.company_id.logo_web, bytes) else order.company_id.logo_web
                logo_html = '<img src="data:image/png;base64,%s" style="max-height:80px;"/>' % logo_b64
            imgs = ''
            for att in self.quote_image_ids:
                if att.datas:
                    imgs += '<div style="margin-bottom:20px;"><img src="data:image/png;base64,%s" style="max-width:400px;max-height:400px;"/></div>' % (att.datas.decode('utf-8') if isinstance(att.datas, bytes) else att.datas)
            if imgs:
                img_html = '<div style="page-break-before:always;"><div style="text-align:right;">%s</div><h2>Product Images</h2>%s</div>' % (logo_html, imgs)

        self.document_html = Markup(str(base_html) + tech_html + img_html)
        # Note: base_html already contains Terms & Conditions at end

    def action_add_images(self):
        self.ensure_one()
        return {
            'type': 'ir.actions.act_window',
            'name': 'Quote Images',
            'res_model': 'sale.quote.preview.wizard',
            'res_id': self.id,
            'view_mode': 'form',
            'view_id': self.env.ref('custom_crm_extended.view_quote_images_form').id,
            'target': 'new',
        }

    def action_download_docx(self):
        self.ensure_one()
        import io, base64
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from odoo.tools import html2plaintext

        doc = Document()

        # Page margins
        from docx.oxml import OxmlElement
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        order = self.order_id
        gst_on = order.x_gst_included

        # Logo right aligned
        if order.company_id.logo_web:
            logo_data = order.company_id.logo_web
            if isinstance(logo_data, str):
                logo_data = logo_data.encode()
            logo_buf = io.BytesIO(base64.b64decode(logo_data))
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = logo_para.add_run()
            run.add_picture(logo_buf, width=Inches(1.5))

        # Quotation No and Date - shaded box like PDF
        qd_table = doc.add_table(rows=1, cols=2)
        qd_table.style = 'Table Grid'
        # Left cell - Quotation No
        left = qd_table.rows[0].cells[0]
        left_para = left.paragraphs[0]
        left_run = left_para.add_run('Quotation No: %s' % (self.quote_name or ''))
        left_run.bold = True
        # Right cell - Date
        right = qd_table.rows[0].cells[1]
        right_para = right.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_run = right_para.add_run('Date: %s' % str(self.quote_date or ''))
        right_run.bold = True
        # Add grey background to both cells
        for cell in [left, right]:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'EAF0FB')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
        doc.add_paragraph('')

        # To
        to_para = doc.add_paragraph()
        to_para.add_run('To, ').bold = True
        to_para.add_run(self.buyer_name or '')

        # Email and Phone
        if order.partner_id.email:
            doc.add_paragraph(order.partner_id.email)
        if order.partner_id.phone:
            doc.add_paragraph(order.partner_id.phone)

        # Subject
        subj_para = doc.add_paragraph()
        subj_para.add_run('Subject: ').bold = True
        subj_para.add_run('Quotation for Products / Services')
        doc.add_paragraph('')
        doc.add_paragraph('Dear Sir,')
        best_offer = self.best_offer_for or ''
        doc.add_paragraph('With reference to your discussion with the undersigned as regards your subject requirement, we are pleased to quote\nour best offer for %s.' % best_offer)
        doc.add_paragraph('')

        # Table
        headers = ['SR No.', 'Item Description', 'HSN', 'Unit Price', 'Discount', 'After Discount', 'Qty', 'Amount']

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Grey background
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'F0F0F0')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)

        for idx, line in enumerate(order.order_line.filtered(lambda l: not l.display_type), 1):
            row_cells = table.add_row().cells
            desc = line.x_product_name or line.product_id.name or ''
            if line.x_make:
                desc += '\nMake: ' + line.x_make
            if hasattr(line, 'x_notes') and line.x_notes:
                desc += '\nNote: ' + line.x_notes
            hsn = line.product_id.l10n_in_hsn_code or ''
            unit_price = line.price_unit or 0
            disc_pct = line.discount or 0
            disc_amt = unit_price * disc_pct / 100
            after_disc = unit_price - disc_amt
            qty = line.product_uom_qty or 0
            amount = line.price_subtotal or 0
            disc_str = '(%s%%)=%s' % (int(disc_pct), int(disc_amt)) if disc_pct else '-'
            row_data = [str(idx), desc, hsn, str(int(unit_price)), disc_str, str(int(after_disc)), str(int(qty)), str(int(amount))]

            for i, val in enumerate(row_data):
                row_cells[i].text = val
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT

        # Totals
        doc.add_paragraph('')
        untax_p = doc.add_paragraph('Gross Total Amount: %s' % int(order.amount_untaxed))
        untax_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if gst_on:
            tax_p = doc.add_paragraph('Tax: %s' % int(order.amount_tax))
            tax_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_p = doc.add_paragraph('Total: %s' % int(order.amount_total))
        total_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_p.runs[0].bold = True

        # Technical Specs - Page 2
        if self.technical_specs_html:
            doc.add_page_break()
            # Logo on page 2
            if order.company_id.logo_web:
                logo_buf2 = io.BytesIO(base64.b64decode(
                    order.company_id.logo_web if isinstance(order.company_id.logo_web, bytes)
                    else order.company_id.logo_web.encode()
                ))
                logo_p2 = doc.add_paragraph()
                logo_p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                logo_p2.add_run().add_picture(logo_buf2, width=Inches(1.5))
            doc.add_heading('Technical Specifications', 1)
            # Parse HTML and render BOTH text and tables in order
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(str(self.technical_specs_html), 'html.parser')
                # Process all top-level elements in order
                for element in soup.children:
                    if hasattr(element, 'name') and element.name == 'table':
                        # Render table
                        rows = element.find_all('tr')
                        if not rows:
                            continue
                        max_cols = max(len(r.find_all(['td','th'])) for r in rows)
                        if max_cols == 0:
                            continue
                        docx_table = doc.add_table(rows=0, cols=max_cols)
                        docx_table.style = 'Table Grid'
                        for row in rows:
                            cells = row.find_all(['td', 'th'])
                            row_cells = docx_table.add_row().cells
                            for i, cell in enumerate(cells):
                                if i < max_cols:
                                    row_cells[i].text = cell.get_text(strip=True)
                                    if cell.name == 'th':
                                        for para in row_cells[i].paragraphs:
                                            for run in para.runs:
                                                run.bold = True
                                        tc = row_cells[i]._tc
                                        tcPr = tc.get_or_add_tcPr()
                                        shd = OxmlElement('w:shd')
                                        shd.set(qn('w:fill'), 'F0F0F0')
                                        shd.set(qn('w:color'), 'auto')
                                        shd.set(qn('w:val'), 'clear')
                                        tcPr.append(shd)
                        doc.add_paragraph('')
                    elif hasattr(element, 'name') and element.name in ['p', 'div']:
                        # Check if contains nested table
                        nested_tables = element.find_all('table')
                        if nested_tables:
                            # Render any text before table
                            pre_text = element.get_text(separator='\n').strip()
                            for nested_table in nested_tables:
                                pre_text = pre_text.replace(nested_table.get_text(separator='\n'), '').strip()
                            if pre_text:
                                doc.add_paragraph(pre_text)
                            # Render nested tables
                            for nested_table in nested_tables:
                                rows = nested_table.find_all('tr')
                                if not rows:
                                    continue
                                max_cols = max(len(r.find_all(['td','th'])) for r in rows)
                                if max_cols == 0:
                                    continue
                                docx_table = doc.add_table(rows=0, cols=max_cols)
                                docx_table.style = 'Table Grid'
                                for row in rows:
                                    cells = row.find_all(['td', 'th'])
                                    row_cells = docx_table.add_row().cells
                                    for i, cell in enumerate(cells):
                                        if i < max_cols:
                                            row_cells[i].text = cell.get_text(strip=True)
                                            if cell.name == 'th':
                                                for para in row_cells[i].paragraphs:
                                                    for run in para.runs:
                                                        run.bold = True
                                                tc = row_cells[i]._tc
                                                tcPr = tc.get_or_add_tcPr()
                                                shd = OxmlElement('w:shd')
                                                shd.set(qn('w:fill'), 'F0F0F0')
                                                shd.set(qn('w:color'), 'auto')
                                                shd.set(qn('w:val'), 'clear')
                                                tcPr.append(shd)
                                doc.add_paragraph('')
                        else:
                            # Plain text paragraph
                            text = element.get_text(strip=True)
                            if text:
                                para = doc.add_paragraph(text)
                                if element.name in ['h1','h2','h3','h4']:
                                    para.runs[0].bold = True
                    elif hasattr(element, 'name') and element.name in ['h1','h2','h3','h4']:
                        text = element.get_text(strip=True)
                        if text:
                            p = doc.add_paragraph(text)
                            p.runs[0].bold = True
                    elif hasattr(element, 'name') and element.name and element.get_text(strip=True):
                        text = element.get_text(strip=True)
                        if text:
                            doc.add_paragraph(text)
            except Exception as e:
                import logging
                logging.getLogger(__name__).error('Tech specs render error: %s', e)
                doc.add_paragraph(html2plaintext(self.technical_specs_html))

        # Images - Page 3
        if self.quote_image_ids:
            doc.add_page_break()
            if order.company_id.logo_web:
                logo_buf3 = io.BytesIO(base64.b64decode(
                    order.company_id.logo_web if isinstance(order.company_id.logo_web, bytes)
                    else order.company_id.logo_web.encode()
                ))
                logo_p3 = doc.add_paragraph()
                logo_p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                logo_p3.add_run().add_picture(logo_buf3, width=Inches(1.5))
            doc.add_heading('Product Images', 1)
            for att in self.quote_image_ids:
                if att.datas:
                    try:
                        img_buf = io.BytesIO(base64.b64decode(att.datas))
                        img_para = doc.add_paragraph()
                        img_para.add_run().add_picture(img_buf, width=Inches(4))
                        pass  # no image name
                    except Exception:
                        pass

        # Terms & Conditions - Last page
        if order.note:
            doc.add_page_break()
            if order.company_id.logo_web:
                logo_buf4 = io.BytesIO(base64.b64decode(
                    order.company_id.logo_web if isinstance(order.company_id.logo_web, bytes)
                    else order.company_id.logo_web.encode()
                ))
                logo_p4 = doc.add_paragraph()
                logo_p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                logo_p4.add_run().add_picture(logo_buf4, width=Inches(1.5))
            doc.add_heading('Terms & Conditions', 2)
            doc.add_paragraph(order.note)

        # Save
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        docx_data = buf.read()

        attachment = self.env['ir.attachment'].create({
            'name': 'Quotation-%s.docx' % (self.quote_name or 'quote'),
            'type': 'binary',
            'datas': base64.b64encode(docx_data).decode(),
            'res_model': 'sale.quote.preview.wizard',
            'res_id': self.id,
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        })
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/content/%s?download=true' % attachment.id,
            'target': 'new',
        }

    def action_send_email(self):
        self.ensure_one()
        import base64
        self._rebuild_document_html()
        order = self.order_id
        best_offer = self.best_offer_for or ''
        subject = 'Quotation %s - %s' % (self.quote_name or '', order.partner_id.name or '')
        body = '<p>Dear Sir,</p><p>' + self.env['ir.config_parameter'].sudo().get_param('sale.quote.intro.template', 'With reference to your discussion with the undersigned as regards your subject requirement, we are pleased to quote our best offer for') + '</p><p>Please find the attached quotation for your reference.</p><p>Best Regards,<br/>%s</p>' % (self.seller_name or '')

        # Generate PDF from HTML content directly
        try:
            from odoo.tools import pdf as pdf_tools
            html = self.document_html or ''
            # Use wkhtmltopdf directly
            import subprocess, tempfile, os
            with tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8') as f:
                f.write('<html><body>%s</body></html>' % str(html))
                html_file = f.name
            pdf_file = html_file.replace('.html', '.pdf')
            subprocess.run(['wkhtmltopdf', '--quiet', html_file, pdf_file], check=True)
            with open(pdf_file, 'rb') as f:
                pdf_content = f.read()
            os.unlink(html_file)
            os.unlink(pdf_file)
            attachment = self.env['ir.attachment'].create({
                'name': 'Quotation-%s.pdf' % (self.quote_name or 'quote'),
                'type': 'binary',
                'datas': base64.b64encode(pdf_content).decode(),
                'res_model': 'sale.order',
                'res_id': order.id,
                'mimetype': 'application/pdf',
            })
            attachment_ids = [(4, attachment.id)]
        except Exception as e:
            import logging
            logging.getLogger(__name__).error('PDF generation failed: %s', e)
            attachment_ids = []

        return {
            'type': 'ir.actions.act_window',
            'name': 'Send Email',
            'res_model': 'mail.compose.message',
            'view_mode': 'form',
            'target': 'new',
            'context': {
                'default_model': 'sale.order',
                'default_res_ids': [order.id],
                'default_subject': subject,
                'default_body': body,
                'default_attachment_ids': attachment_ids,
                'default_partner_ids': [(4, order.partner_id.id)] if order.partner_id else [],
            },
        }

    def action_print_edited_pdf(self):
        self.ensure_one()
        # Always rebuild to include technical specs and images
        self._rebuild_document_html()
        return self.env['ir.actions.report'].search([('report_name','=','custom_crm_extended.report_sale_quote_preview_wizard')], limit=1).report_action(self)
