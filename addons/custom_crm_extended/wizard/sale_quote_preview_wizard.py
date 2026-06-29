# -*- coding: utf-8 -*-

def _indian_format(n):
    """Format number in Indian style: 1,00,000"""
    n = int(n)
    s = str(n)
    if len(s) <= 3:
        return s
    last3 = s[-3:]
    rest = s[:-3]
    parts = []
    while len(rest) > 2:
        parts.append(rest[-2:])
        rest = rest[:-2]
    if rest:
        parts.append(rest)
    parts.reverse()
    return ','.join(parts) + ',' + last3

import re
from odoo import api, fields, models
from odoo.tools import html2plaintext
from markupsafe import Markup


class SaleQuotePreviewWizard(models.TransientModel):
    _name = 'sale.quote.preview.wizard'
    _description = 'Editable Quotation Preview Wizard'

    order_id = fields.Many2one('sale.order', string='Quotation', required=True, readonly=True)
    x_gst_included = fields.Boolean(string='Include GST', default=True)
    seller_name = fields.Char(string='Seller / Company Name')
    buyer_name = fields.Char(string='Party Name')
    contact_person = fields.Char(string='Contact Person')
    contact_function = fields.Char(string='Job Title / Designation')
    quote_name = fields.Char(string='Quotation No.')
    quote_date = fields.Date(string='Quotation Date')
    valid_until = fields.Date(string='Valid Upto')
    subject = fields.Char(string='Subject')
    best_offer_for = fields.Char(string='Best Offer For', help='Text after best offer for in the intro paragraph')

    company_logo = fields.Binary(string='Company Logo')
    document_html = fields.Html(string='Editable Quote Content')
    x_table_html = fields.Html(string='Commercial Table HTML')
    x_terms_html = fields.Html(string='Terms HTML')
    technical_specs_html = fields.Html(string='Technical Specifications')
    selected_term_ids = fields.Many2many(
        'sale.terms.condition',
        'sale_quote_wizard_terms_rel',
        'wizard_id', 'term_id',
        string='Terms & Conditions',
        default=lambda self: self.env['sale.terms.condition'].search([], order='sequence, id')
    )
    quote_image_ids = fields.Many2many(
        'ir.attachment', 'sale_quote_wizard_image_rel',
        'wizard_id', 'attachment_id',
        string='Quote Images'
    )



    terms_html_widget = fields.Html(
        string='Terms Selection',
        store=True,
    )

    @api.depends('selected_term_ids')
    def _compute_terms_html(self):
        for rec in self:
            all_terms = self.env['sale.terms.condition'].search([], order='sequence, id')
            html = '<div class="o_terms_checkboxes">'
            for term in all_terms:
                checked = 'checked' if term in rec.selected_term_ids else ''
                html += '''<div style="margin:4px 0;">
                    <input type="checkbox" data-term-id="%s" %s style="margin-right:8px;"/>
                    <label style="font-size:13px;">%s</label>
                </div>''' % (term.id, checked, term.name)
            html += '</div>'
            rec.terms_html_widget = html





    def default_get(self, fields_list):
        res = super().default_get(fields_list)
        order = self.env['sale.order'].browse(self.env.context.get('default_order_id'))
        if not order:
            return res
        # Restore draft term selection if saved, else default to all terms
        all_terms = self.env['sale.terms.condition'].search([], order='sequence, id')
        # has_draft: True if any draft field has been saved before
        has_draft = bool(order.x_draft_term_ids or order.x_draft_quote_name or order.x_draft_best_offer or order.x_draft_tech_specs)
        import logging
        _log = logging.getLogger(__name__)
        _log.warning('DRAFT DEBUG: order=%s has_draft=%s x_draft_quote_name=%s x_draft_valid_until=%s', order.name, has_draft, order.x_draft_quote_name, order.x_draft_valid_until)
        if order.x_draft_term_ids:
            res['selected_term_ids'] = [(6, 0, order.x_draft_term_ids.ids)]
        elif all_terms:
            res['selected_term_ids'] = [[6, 0, all_terms.ids]]
        # Use saved draft GST if a draft exists, else fall back to order setting
        gst_on = order.x_draft_gst_included if has_draft else order.x_gst_included

        rows = ''
        for idx, line in enumerate(order.order_line.filtered(lambda l: not l.display_type), 1):
            part_no = line.x_product_code or line.product_id.default_code or ''
            note = (line.x_notes if hasattr(line, 'x_notes') and line.x_notes else '')
            hsn = line.product_id.l10n_in_hsn_code or ''
            make = line.x_make or ''
            unit_price = line.price_unit or 0
            discount_pct = (line.discount or 0)
            disc_amount = unit_price * discount_pct / 100
            after_discount = unit_price - disc_amount
            qty = line.product_uom_qty or 0
            amount = line.price_subtotal or 0

            # Tax
            tax_str = ''
            for tax in line.tax_ids:
                tname = (tax.name or '').upper()
                tax_str += '%s(%s%%) ' % (tname, int(tax.amount))
            tax_str = tax_str.strip() or '-'

            # Discount
            if discount_pct:
                disc_str = '(%s%%)=%s' % (int(discount_pct), int(disc_amount)) if discount_pct else '-'
            else:
                disc_str = '-'

            # Description
            desc_html = line.x_product_name or line.product_id.name or ''
            if part_no:
                desc_html += '<br/><small style="color:#666;">Part No: %s</small>' % part_no
            if make:
                desc_html += '<br/><small style="color:#555;">Make: %s</small>' % make
            if note:
                desc_html += '<br/><small style="color:#888;font-style:italic;">Note: %s</small>' % note

            # Image
            img_html = ''
            if line.product_id.image_128:
                img_b64 = line.product_id.image_128.decode('utf-8') if isinstance(line.product_id.image_128, bytes) else line.product_id.image_128
                img_html = ""

            tax_cell = '<td style="text-align:center;" contenteditable="true">%s</td>' % tax_str if gst_on else ''

            rows += (
                '<tr>'
                '<td style="text-align:center;white-space:nowrap;">%s</td>'
                '<td>%s%s</td>'
                '<td style="text-align:center;white-space:nowrap;">%s</td>'
                '<td style="text-align:right;" contenteditable="false">%s</td>'
                '<td style="text-align:right;" contenteditable="false">%s</td>'
                '<td style="text-align:right;" contenteditable="false">%s</td>'
                '<td style="text-align:right;" contenteditable="false">%s</td>'
                '<td style="text-align:right;" contenteditable="false">%s</td>'
                '</tr>'
            ) % (idx, desc_html, img_html, hsn, int(unit_price), disc_str, int(after_discount), int(qty), int(amount))

        # Technical Specs Page
        tech_rows = ''
        for idx, line in enumerate(order.order_line.filtered(lambda l: not l.display_type), 1):
            if line.x_technical_specs:
                pname = line.x_product_name or line.product_id.name or ''
                tech_rows += (
                    '<h4>%s. %s</h4>'
                    '<p>%s</p>'
                ) % (idx, pname, line.x_technical_specs.replace('\n', '<br/>'))

        # Images Page
        img_rows = ''
        for idx, line in enumerate(order.order_line.filtered(lambda l: not l.display_type), 1):
            img_src = ''
            if line.x_product_image:
                img_b64 = line.x_product_image.decode('utf-8') if isinstance(line.x_product_image, bytes) else line.x_product_image
                img_src = '<img src="data:image/png;base64,%s" style="max-width:300px;max-height:300px;"/>' % img_b64
            elif line.product_id.image_1920:
                img_b64 = line.product_id.image_1920.decode('utf-8') if isinstance(line.product_id.image_1920, bytes) else line.product_id.image_1920
                img_src = '<img src="data:image/png;base64,%s" style="max-width:300px;max-height:300px;"/>' % img_b64
            if img_src:
                pname = line.x_product_name or line.product_id.name or ''
                img_rows += '<div style="margin-bottom:20px;"><h4>%s</h4>%s</div>' % (pname, img_src)

        # Get product categories from linked opportunity only
        product_cats = self.env.context.get('best_offer_for', '')
        if not product_cats:
            if order.opportunity_id and order.opportunity_id.x_product_category_ids:
                product_cats = ', '.join(order.opportunity_id.x_product_category_ids.mapped('name'))

        tax_th = '<th style="text-align:center;">Tax</th>' if gst_on else ''
        tax_amount_row = ''

        # Logo
        logo_html = ''
        if order.company_id.logo_web:
            logo_b64 = order.company_id.logo_web.decode('utf-8') if isinstance(order.company_id.logo_web, bytes) else order.company_id.logo_web
            logo_html = '<img src="data:image/png;base64,%s" style="max-height:80px;"/>' % logo_b64

        intro_text = self.env['ir.config_parameter'].sudo().get_param('sale.quote.intro.template', 'With reference to your discussion with the undersigned as regards your subject requirement, we are pleased to quote our best offer for')
        # INTRO section (header info, before table)
        intro_html = (
            '<p style="font-size:11px;margin-top:10px;"><b>To,</b></p>'
            '<p style="font-size:11px;margin:2px 0;">%s</p>'
            '<p style="margin:2px 0;">%s</p>'
            '<p style="margin:2px 0;">%s</p>'
            '<p style="margin:2px 0;">%s</p>'
            '<br/>'
            '<p style="font-size:11px;margin-bottom:10px;"><b>Subject:</b> Quotation for Products / Services</p>'
            '<br/>'
            '<p style="font-size:11px;">Dear Sir,</p>'
            '<p>%s</p>'
        ) % (
            order.partner_id.name or '',
            order.partner_id.city or '',
            order.partner_id.email or '',
            order.partner_id.phone or '',
            intro_text,
        )

        # TABLE section (commercial table + totals) - keep on its own page
        table_html = (
            '<div style="' + ('page-break-before:always;' if (bool(re.sub(r'<[^>]+>', '', str(self.technical_specs_html or '')).strip()) or bool(self.quote_image_ids)) else '') + 'padding:40px 0 10px 0;">'
            '<p style="text-align:center;font-size:15px;font-weight:bold;margin:0 0 12px 0;padding:8px 0;">Quotation</p>'
            '<table border="1" cellpadding="6" cellspacing="0" style="width:100%%;border-collapse:collapse;font-size:11px;page-break-inside:avoid;" contenteditable="false">'
            '<thead><tr style="background:#f0f0f0;">'
            '<th style="text-align:center;width:55px;white-space:nowrap;">SR No.</th>'
            '<th style="text-align:left;">Item Description</th>'
            '<th style="text-align:center;">HSN</th>'
            '<th style="text-align:right;">Unit Price</th>'
            '<th style="text-align:right;">Discount</th>'
            '<th style="text-align:right;">After Discount</th>'
            '<th style="text-align:right;">Qty</th>'
            '<th style="text-align:right;">Amount</th>'
            '</tr></thead>'
            '<tbody>%s'
            '<tr><td colspan="5" style="text-align:right;border:none;"><b>Gross Total Amount INR:</b></td>'
            '<td style="text-align:right;border:1px solid #000;"><b>%s</b></td></tr>'
            '%s'
            '<tr><td colspan="5" style="text-align:right;border:none;"><b>Net Total Amount INR:</b></td>'
            '<td style="text-align:right;border:1px solid #000;"><b>%s</b></td></tr>'
            '</tbody>'
            '</table>'
            '</div>'
        ) % (
            rows,
            _indian_format(order.amount_untaxed),
            tax_amount_row,
            _indian_format(order.amount_total),
        )

        # TERMS section
        terms_content = str(order.note or '')
        # Footer: selected company's address
        cp = order.company_id.partner_id
        addr_parts = [p for p in [cp.street, cp.street2, cp.city,
                                   cp.state_id.name if cp.state_id else '',
                                   cp.zip, cp.country_id.name if cp.country_id else ''] if p]
        addr_line = ', '.join(addr_parts)
        contact_parts = []
        if cp.phone:
            contact_parts.append('Phone: %s' % cp.phone)
        if cp.email:
            contact_parts.append('Email: %s' % cp.email)
        contact_line = ' | '.join(contact_parts)
        footer_html = ''  # Address shown in sidebar only
        terms_html = (
            '<div style="margin-top:20px;">'
            '<h3 style="text-align:center;font-size:11px;font-weight:bold;margin:12px 0 10px 0;">Terms &amp; Conditions</h3>'
            '<div style="font-size:11px;line-height:1.6;">%s</div>'
            '%s'
            '</div>'
        ) % (terms_content, footer_html)

        html = intro_html  # default_get only needs intro for initial preview

        _unused = (
            rows, _indian_format(order.amount_untaxed), tax_amount_row, _indian_format(order.amount_total), order.note or ''
        )

        # Build tech specs section
        tech_section = ''
        if tech_rows:
            tech_section = (
                '<div style="margin-top:20px;">'
                '<div style="text-align:right;">%s</div>'
                '<h2>Technical Specifications</h2>%s'
                '</div>'
            ) % (logo_html, tech_rows)

        # Build images section
        img_section = ''
        if img_rows:
            img_section = (
                '<div style="margin-top:20px;">'
                '<div style="text-align:right;">%s</div>'
                '%s'
                '</div>'
            ) % (logo_html, img_rows)

        res.update({
            'order_id': order.id,
            'x_gst_included': gst_on,
            'seller_name': order.company_id.name or '',
            'buyer_name': order.x_draft_buyer_name or order.partner_id.name or '',
            'contact_person': order.x_draft_contact_person or order.x_contact_person or (order.opportunity_id.contact_name if order.opportunity_id else '') or '',
            'contact_function': order.x_draft_contact_function or (order.opportunity_id.function if order.opportunity_id else ''),
            'quote_name': order.x_draft_quote_name or order.name or '',
            'quote_date': order.x_draft_quote_date or (order.date_order.date() if order.date_order else fields.Date.today()),
            'valid_until': order.x_draft_valid_until or order.validity_date,
            'subject': order.x_draft_subject or 'Quotation for Products / Services',
            'best_offer_for': order.x_draft_best_offer or product_cats,
            'company_logo': order.company_id.logo_web,
            'document_html': Markup(html),
            'x_table_html': table_html,
            'x_terms_html': terms_html,
            'technical_specs_html': getattr(order, 'x_draft_tech_specs', None) or False,
            'quote_image_ids': [(6, 0, getattr(order, 'x_draft_image_ids', self.env['ir.attachment']).ids)],
        })
        return res

    def action_save_draft(self):
        self.ensure_one()
        order = self.order_id
        if not order:
            return
        # Save selected terms to order note
        if self.selected_term_ids:
            items = ''.join('<li>%s</li>' % (t.content or '') for t in self.selected_term_ids.sorted('sequence'))
            order.sudo().write({'note': '<ol>%s</ol>' % items})
        vals = {
            'x_draft_tech_specs': self.technical_specs_html,
            'x_draft_best_offer': self.best_offer_for,
            'x_draft_gst_included': self.x_gst_included,
            'x_draft_valid_until': self.valid_until,
            'x_draft_term_ids': [(6, 0, self.selected_term_ids.ids)],
            'x_draft_quote_name': self.quote_name,
            'x_draft_contact_person': self.contact_person,
            'x_draft_contact_function': self.contact_function,
            'x_draft_subject': self.subject,
            'x_draft_quote_date': self.quote_date,
            'x_draft_buyer_name': self.buyer_name,
        }
        if hasattr(order, 'x_draft_image_ids'):
            vals['x_draft_image_ids'] = [(6, 0, self.quote_image_ids.ids)]
        order.sudo().write(vals)
        # Reopen wizard keeping all state + show notification
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'sale.quote.preview.wizard',
            'res_id': self.id,
            'view_mode': 'form',
            'target': 'new',
            'flags': {'mode': 'edit'},
            'context': {
                'default_order_id': order.id,
                'show_notification': 'Draft Saved!',
            },
        }

    def action_add_technical_specs(self):
        self.ensure_one()
        return {
            'type': 'ir.actions.act_window',
            'name': 'Technical Specifications',
            'res_model': 'sale.quote.preview.wizard',
            'res_id': self.id,
            'view_mode': 'form',
            'view_id': self.env.ref('custom_crm_extended.view_quote_tech_specs_form').id,
            'target': 'new',
        }

    def action_save_and_reopen(self):
        self.ensure_one()
        self._rebuild_document_html()
        # Reopen the main preview wizard keeping state
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'sale.quote.preview.wizard',
            'res_id': self.id,
            'view_mode': 'form',
            'target': 'new',
            'flags': {'mode': 'edit'},
        }


    def _style_html_tables(self, html_content):
        """Add proper border and styling to all tables in HTML content"""
        from markupsafe import Markup
        import re
        html_str = str(html_content)
        # Style table tags
        html_str = re.sub(
            r'<table(?![^>]*style)[^>]*>',
            '<table style="width:100%;border-collapse:collapse;font-size:11px;margin-bottom:10px;" border="1" cellpadding="6" cellspacing="0">',
            html_str
        )
        # Style th tags
        html_str = re.sub(
            r'<th(?![^>]*style)[^>]*>',
            '<th style="background:#f0f0f0;border:1px solid #ccc;padding:6px;text-align:left;font-weight:bold;">',
            html_str
        )
        # Style td tags
        html_str = re.sub(
            r'<td(?![^>]*style)[^>]*>',
            '<td style="border:1px solid #ccc;padding:6px;text-align:left;">',
            html_str
        )
        return Markup(html_str)

    def _rebuild_document_html(self):
        from markupsafe import Markup
        order = self.order_id
        if not order:
            return

        comp = order.company_id
        p = order.partner_id

        # ── LOGO ──
        logo_html = ''
        if comp.logo_web:
            b64 = comp.logo_web.decode('utf-8') if isinstance(comp.logo_web, bytes) else comp.logo_web
            logo_html = '<img src="data:image/png;base64,%s" style="max-height:70px;max-width:200px;object-fit:contain;"/>' % b64

        # ── PRODUCT ROWS ──
        gst_on = self.x_gst_included if self.x_gst_included is not None else order.x_gst_included
        rows = ''
        for idx, line in enumerate(order.order_line.filtered(lambda l: not l.display_type), 1):
            part_no = line.x_product_code or line.product_id.default_code or ''
            note = line.x_notes if hasattr(line, 'x_notes') and line.x_notes else ''
            hsn = line.product_id.l10n_in_hsn_code or ''
            make = line.x_make or ''
            unit_price = line.price_unit or 0
            discount_pct = line.discount or 0
            disc_amount = unit_price * discount_pct / 100
            after_disc = unit_price - disc_amount
            qty = line.product_uom_qty or 0
            amount = line.price_subtotal or 0
            disc_str = '(%s%%)=%s' % (int(discount_pct), int(disc_amount)) if discount_pct else '-'
            desc = line.x_product_name or line.product_id.name or ''
            if make: desc += '<br/><small style="color:#666;">Make: %s</small>' % make
            if part_no: desc += '<br/><small style="color:#888;">Part No: %s</small>' % part_no
            if note: desc += '<br/><small style="color:#999;font-style:italic;">%s</small>' % note
            rows += (
                '<tr style="background:%s;">' % ('#f9f9f9' if idx % 2 == 0 else '#fff')
                + '<td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td>' % idx
                + '<td style="padding:6px 8px;border:1px solid #ddd;">%s</td>' % desc
                + '<td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td>' % hsn
                + '<td style="text-align:right;padding:6px 8px;border:1px solid #ddd;">%s</td>' % int(unit_price)
                + '<td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td>' % disc_str
                + '<td style="text-align:right;padding:6px 8px;border:1px solid #ddd;">%s</td>' % int(after_disc)
                + '<td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td>' % int(qty)
                + '<td style="text-align:right;padding:6px 8px;border:1px solid #ddd;">%s</td>' % int(amount)
                + '</tr>'
            )

        # ── TAX ROW ──
        tax_row = ''
        if gst_on and order.amount_tax:
            tax_row = '<p style="text-align:right;margin:4px 0;font-size:11px;">Tax (GST): <b>%s</b></p>' % int(order.amount_tax)

        # ── SUBJECT / INTRO ──
        subject = self.subject or 'Quotation for Products / Services'
        best_offer = self.best_offer_for or ''
        intro_text = self.env["ir.config_parameter"].sudo().get_param(
            "sale.quote.intro.template",
            "With reference to your discussion with the undersigned as regards your subject requirement, we are pleased to quote our best offer for"
        )

        # ── ADDRESS LINES ──
        addr_parts = [x for x in [
            p.street or '',
            p.street2 or '',
            ('%s %s' % (p.city or '', p.zip or '')).strip(),
            p.state_id.name if p.state_id else '',
        ] if x]
        addr_html = ''.join('<p style="margin:0 0 1px 0;font-size:11px;">%s</p>' % a for a in addr_parts)
        addr_html += '<br/>'

        if p.phone:
            addr_html += '<p style="margin:0 0 1px 0;font-size:11px;">Ph: %s</p>' % p.phone
        if p.email:
            addr_html += '<p style="margin:0 0 1px 0;font-size:11px;">Email: %s</p>' % p.email

        # ── PAGE 1: INTRO ──
        _pdf_date = order.date_order.date() if order.date_order else fields.Date.today()
        _pdf_date_str = _pdf_date.strftime('%d-%m-%Y')
        intro_html = (
            '<div style="font-family:Calibri,sans-serif;font-size:11px;line-height:1.6;color:#222;margin-top:0;">'
            '<div style="overflow:hidden;margin-bottom:12px;">'
            '<span style="float:left;font-weight:bold;">Quotation No: %s</span>'
            '<span style="float:right;font-weight:bold;">Date: %s</span>'
            '</div>'
            '<div style="clear:both;"></div>'
            ) % (self.quote_name or order.name or '', _pdf_date_str)
        intro_html += (
            '<p style="margin:0 0 3px 0;"><b>To,</b></p>'
            + ('<p style="margin:0 0 1px 0;">%s</p>' % self.contact_person if self.contact_person else '')
            + '<p style="margin:0 0 1px 0;font-weight:bold;">%s</p>' % (p.name or '')

            + addr_html
            + '<br/>'
            + '<p style="margin:6px 0;"><b>Subject:</b> %s</p>' % subject
            + '<br/>'
            + '<p style="margin:4px 0;">Dear Sir,</p>'
            + '<p style="margin:6px 0;">%s%s</p>' % (intro_text, (' ' + best_offer + '.') if best_offer else '.')
            + '</div>'
        )
        # ── TECH SPECS ──
        tech_html = ''
        # Check if technical specs has actual content (not just empty HTML)
        _tech_text = self.technical_specs_html or ''
        import re as _re
        _tech_clean = _re.sub(r'<[^>]+>', '', str(_tech_text)).strip()
        if self.technical_specs_html and _tech_clean:
            from markupsafe import Markup as _M
            styled = self._style_html_tables(self.technical_specs_html)
            tech_html = (
                '<div style="page-break-before:always;margin-top:16px;font-family:Calibri,sans-serif;font-size:11px;">'
                '<p style="font-weight:bold;font-size:11px;margin:10px 0 16px 0;">Technical Specifications</p><br/>'
                + str(styled)
                + '</div>'
            )

        # ── IMAGES ──
        img_html = ''
        if self.quote_image_ids:
            imgs = ''
            for att in self.quote_image_ids:
                if att.datas:
                    b64 = att.datas.decode('utf-8') if isinstance(att.datas, bytes) else att.datas
                    imgs += (
                        '<div style="display:inline-block;width:47%%;margin:1%%;vertical-align:top;text-align:center;">'
                        '<img src="data:image/png;base64,%s" style="max-width:100%%;max-height:220px;border:1px solid #ddd;padding:3px;"/>'
                        '</div>'
                    ) % b64
            if imgs:
                img_html = '<div style="margin-top:12px;">%s</div>' % imgs

        # ── CHECK DISCOUNT FOR PDF TABLE ──
        order_lines_pdf = order.order_line
        has_discount_pdf = any(l.discount for l in order_lines_pdf)
        has_overall_disc = getattr(order, 'x_flat_discount_pct', 0) or 0

        # Build PDF rows
        rows = ''
        idx2 = 0
        # Collect notes per product line
        note_map = {}
        prev_product_line = None
        for line in order_lines_pdf:
            if line.display_type in ('line_note', 'note'):
                if prev_product_line is not None:
                    note_map.setdefault(prev_product_line, []).append(line.name or '')
            else:
                prev_product_line = line.id

        for line in order_lines_pdf:
            if line.display_type:
                continue
            idx2 += 1
            part_no = line.x_product_code or line.product_id.default_code or ''
            note = line.x_notes if hasattr(line, 'x_notes') and line.x_notes else ''
            hsn = line.product_id.l10n_in_hsn_code or ''
            make = line.x_make or ''
            unit_price = line.price_unit or 0
            discount_pct = (line.discount or 0)
            qty = int(line.product_uom_qty or 0)
            amount = line.price_subtotal or 0
            desc = line.x_product_name or line.product_id.product_tmpl_id.with_context(lang='en_US').name or ''
            if make: desc += '<br/><b>Make:</b> %s' % make
            if note: desc += '<br/><b>Description:</b> %s' % note
            for n in note_map.get(line.id, []):
                desc += '<br/><b>Description:</b> <i style="color:#333;">%s</i>' % n
            row_bg = '#f9f9f9' if idx2 % 2 == 0 else '#fff'
            if has_discount_pdf:
                disc_str = '%s%%' % int(discount_pct) if discount_pct else '-'
                rows += ('<tr style="background:%s;"><td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td><td style="padding:6px 8px;border:1px solid #ddd;">%s</td><td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td><td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td><td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td><td style="text-align:right;padding:6px 8px;border:1px solid #ddd;">%s</td><td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td><td style="text-align:right;padding:6px 8px;border:1px solid #ddd;">%s</td></tr>') % (row_bg, idx2, desc, part_no, hsn, qty, _indian_format(unit_price), disc_str, _indian_format(amount))
            else:
                rows += ('<tr style="background:%s;"><td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td><td style="padding:6px 8px;border:1px solid #ddd;">%s</td><td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td><td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td><td style="text-align:center;padding:6px 4px;border:1px solid #ddd;">%s</td><td style="text-align:right;padding:6px 8px;border:1px solid #ddd;">%s</td><td style="text-align:right;padding:6px 8px;border:1px solid #ddd;">%s</td></tr>') % (row_bg, idx2, desc, part_no, hsn, qty, _indian_format(unit_price), _indian_format(amount))

        # ── PAGE 2: QUOTATION TABLE (new page) ──
        # Build headers based on discount
        if has_discount_pdf:
            th_html = (
                '<th style="padding:8px 5px;text-align:center;border:1px solid #2c3e50;width:40px;">SR No.</th>'
                '<th style="padding:8px;text-align:left;border:1px solid #2c3e50;">Item Description</th>'
                '<th style="padding:8px;text-align:center;border:1px solid #2c3e50;">Part No</th>'
                '<th style="padding:8px;text-align:center;border:1px solid #2c3e50;">HSN</th>'
                '<th style="padding:8px;text-align:center;border:1px solid #2c3e50;">Qty</th>'
                '<th style="padding:8px;text-align:right;border:1px solid #2c3e50;">Unit Price</th>'
                '<th style="padding:8px;text-align:center;border:1px solid #2c3e50;">Discount %</th>'
                '<th style="padding:8px;text-align:right;border:1px solid #2c3e50;">Amount</th>'
            )
        else:
            th_html = (
                '<th style="padding:8px 5px;text-align:center;border:1px solid #2c3e50;width:40px;">SR No.</th>'
                '<th style="padding:8px;text-align:left;border:1px solid #2c3e50;">Item Description</th>'
                '<th style="padding:8px;text-align:center;border:1px solid #2c3e50;">Part No</th>'
                '<th style="padding:8px;text-align:center;border:1px solid #2c3e50;">HSN</th>'
                '<th style="padding:8px;text-align:center;border:1px solid #2c3e50;">Qty</th>'
                '<th style="padding:8px;text-align:right;border:1px solid #2c3e50;">Unit Price</th>'
                '<th style="padding:8px;text-align:right;border:1px solid #2c3e50;">Amount</th>'
            )

        # Overall discount totals
        totals_html = '<p style="margin:4px 0;">Gross Total Amount INR: <b>%s</b></p>' % int(order.amount_untaxed)
        if has_overall_disc:
            disc_overall = order.amount_untaxed * has_overall_disc / 100
            totals_html += '<p style="margin:4px 0;">Discount (%s%%): %s</p>' % (int(has_overall_disc), int(disc_overall))
        net = order.amount_untaxed - (order.amount_untaxed * has_overall_disc / 100) if has_overall_disc else order.amount_untaxed
        totals_html += '<p style="margin:4px 0;font-size:14px;font-weight:bold;border-top:2px solid #333;padding-top:6px;">Net Total Amount INR: %s</p>' % int(net)

        col = 7 if has_discount_pdf else 6
        original_amount = sum(l.price_unit * l.product_uom_qty for l in order.order_line.filtered(lambda x: not x.display_type))
        untaxed = order.amount_untaxed
        overall_disc_pct = getattr(order, 'x_flat_discount_pct', 0) or 0
        overall_disc_amt = untaxed * overall_disc_pct / 100 if overall_disc_pct else 0
        net = untaxed - overall_disc_amt

        # Get unique tax rates
        tax_rates = set()
        for line in order.order_line.filtered(lambda x: not x.display_type):
            for tax in line.tax_ids:
                tax_rates.add(float(tax.amount))
        total_tax_rate = sum(tax_rates)

        # Grand total = net + GST on net
        grand_total = net + (net * total_tax_rate / 100) if (gst_on and tax_rates) else net

        # Totals below table
        totals_below = '<div style="margin-top:8px;text-align:right;font-size:11px;font-family:Calibri,sans-serif;">'
        totals_below += '<p style="margin:3px 0;">Total Amount: <b>%s</b></p>' % _indian_format(untaxed)
        if overall_disc_pct:
            totals_below += '<p style="margin:3px 0;">Overall Discount (%s%%): <b>%s</b></p>' % (int(overall_disc_pct), _indian_format(overall_disc_amt))
        if gst_on and tax_rates:
            totals_below += '<p style="margin:3px 0;">GST (%s%%)</p>' % int(total_tax_rate)
        final_label = 'Final Discount Amount:' if overall_disc_pct else 'Final Total Amount:'
        totals_below += '<p style="margin:3px 0;font-size:13px;font-weight:bold;">%s %s</p>' % (final_label, _indian_format(grand_total))
        totals_below += '</div>'

        table_html = (
            '<div style="' + ('page-break-before:always' if (bool(re.sub(r'<[^>]+>', '', str(self.technical_specs_html or '')).strip()) or bool(self.quote_image_ids)) else 'margin-top:30px') + ';font-family:Calibri,sans-serif;">'
            '<p style="text-align:center;font-size:15px;font-weight:bold;margin:16px 0 14px 0;letter-spacing:2px;color:#2c3e50;">QUOTATION</p>'
            '<table style="width:100%%;border-collapse:collapse;font-size:11px;">'
            '<thead><tr style="background:#2c3e50;color:#fff;">%s</tr></thead>'
            '<tbody>%s</tbody>'
            '</table>'
            '%s'
            '<div style="display:none;">%s</div>'
            '</div>'
        )
        table_html = table_html % (th_html, rows, totals_below, totals_html)

        # ── TERMS ──
        terms_html = ''
        if self.selected_term_ids:
            items = ''.join('<li>%s</li>' % (t.content or '') for t in self.selected_term_ids.sorted('sequence'))
            note_content = '<ol>%s</ol>' % items
        elif order.note:
            note_content = str(order.note)
        else:
            note_content = ''
        if note_content:
            terms_html = (
                '<div style="' + ('page-break-before:always;padding-top:20px;' if not (bool(re.sub(r'<[^>]+>', '', str(self.technical_specs_html or '')).strip()) or bool(self.quote_image_ids)) else 'margin-top:30px;') + 'font-family:Calibri,sans-serif;font-size:11px;">'
                '<p style="text-align:center;font-weight:bold;font-size:11px;'
                'padding-bottom:6px;margin-bottom:10px;">'
                'Terms &amp; Conditions</p>'
                + note_content
                + '</div>'
            )

        # ── COMBINE ──
        # Closing signature with Regards
        user = order.user_id
        user_name = user.name or ''
        user_job = user.function if hasattr(user, 'function') else ''
        user_phone = user.partner_id.phone or ''
        user_email = user.partner_id.email or user.email or ''
        company_name = order.company_id.name or ''
        closing_html = '<div style="margin-top:30px;font-family:Calibri,sans-serif;font-size:11px;">'
        closing_html += '<p style="margin:4px 0;">Thanking You,</p>'
        closing_html += '<p style="margin:4px 0;">Sincerely,</p>'
        closing_html += '<br/>'
        closing_html += '<p style="margin:4px 0;font-weight:bold;">%s</p>' % user_name
        if user_job:
            closing_html += '<p style="margin:4px 0;">%s</p>' % user_job
        closing_html += '<p style="margin:4px 0;">%s</p>' % company_name
        if user_phone:
            closing_html += '<p style="margin:4px 0;">Ph: %s</p>' % user_phone
        if user_email:
            closing_html += '<p style="margin:4px 0;">Email: %s</p>' % user_email
        closing_html += '</div>'

        css_reset = '<style>a{text-decoration:none!important;color:inherit!important;color:#000!important;}a:link,a:visited,a:hover{text-decoration:none!important;border-bottom:none!important;}header,.header,div.header{border:none!important;border-top:none!important;border-bottom:none!important;box-shadow:none!important;}</style>'
        full_html = css_reset + intro_html + img_html + tech_html + table_html + terms_html + closing_html
        self.sudo().write({'document_html': Markup(full_html)})
        import logging
        logging.getLogger(__name__).warning("REBUILD OK - len:%s has_quotation:%s", len(full_html), 'QUOTATION' in full_html)

    def action_apply_terms(self):
        self.ensure_one()
        order = self.order_id
        if self.selected_term_ids:
            items = ''.join('<li>%s</li>' % (t.content or '') for t in self.selected_term_ids.sorted('sequence'))
            order.sudo().write({'note': '<ol>%s</ol>' % items})
        self._rebuild_document_html()
        return True

    def action_apply_terms_and_close(self):
        self.action_apply_terms()
        return {'type': 'ir.actions.act_window_close'}

    def action_save_and_close(self):
        self.ensure_one()
        if self.selected_term_ids:
            items = ''.join('<li>%s</li>' % (t.content or '') for t in self.selected_term_ids.sorted('sequence'))
            self.order_id.sudo().write({'note': '<ol>%s</ol>' % items})
        return {'type': 'ir.actions.act_window_close'}

    def action_add_images(self):
        self.ensure_one()
        return {
            'type': 'ir.actions.act_window',
            'name': 'Quote Images',
            'res_model': 'sale.quote.preview.wizard',
            'res_id': self.id,
            'view_mode': 'form',
            'view_id': self.env.ref('custom_crm_extended.view_quote_images_form').id,
            'target': 'new',
        }

    def action_download_docx(self):
        self.ensure_one()
        # Save selected terms to order note
        if self.selected_term_ids:
            items = ''.join('<li>%s</li>' % (t.content or '') for t in self.selected_term_ids.sorted('sequence'))
            self.order_id.sudo().write({'note': '<ol>%s</ol>' % items})

        import io, base64
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from odoo.tools import html2plaintext

        doc = Document()
        # Set default font Calibri 11 for entire document
        from docx.oxml.ns import qn as _qn_font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.element.rPr.rFonts.set(_qn_font('w:eastAsia'), 'Calibri')
        for s in doc.styles:
            try:
                s.font.name = 'Calibri'
                s.font.size = Pt(11)
            except:
                pass

        # Page margins
        from docx.oxml import OxmlElement
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        order = self.order_id
        gst_on = order.x_gst_included

        # ── Word header repeats on every page ──
        from docx.oxml import OxmlElement as _OE
        from docx.oxml.ns import qn as _qn
        section0 = doc.sections[0]
        section0.different_first_page_header_footer = False
        hdr = section0.header
        hdr.is_linked_to_previous = False
        # Clear existing header paragraphs
        for p in hdr.paragraphs:
            p.clear()
        # Header: Logo only on right
        hdr_para = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if order.company_id.logo_web:
            ld = order.company_id.logo_web
            if isinstance(ld, str): ld = ld.encode()
            lbuf = io.BytesIO(base64.b64decode(ld))
            hdr_para.add_run().add_picture(lbuf, width=Inches(2.5))



        # Quotation No and Date - shaded box like PDF

        # Quotation No and Date before To block
        qd_table = doc.add_table(rows=1, cols=2)
        qd_table.style = 'Table Grid'
        from docx.oxml import OxmlElement as _OE4
        from docx.oxml.ns import qn as _qn4
        tbl_pr2 = qd_table._tbl.tblPr
        tbl_borders2 = _OE4('w:tblBorders')
        for bn2 in ['top','left','bottom','right','insideH','insideV']:
            b2 = _OE4('w:%s' % bn2)
            b2.set(_qn4('w:val'), 'none')
            tbl_borders2.append(b2)
        tbl_pr2.append(tbl_borders2)
        ql = qd_table.rows[0].cells[0].paragraphs[0]
        ql.alignment = WD_ALIGN_PARAGRAPH.LEFT
        qlr = ql.add_run('Quotation No: %s' % (self.quote_name or order.name or ''))
        qlr.bold = True
        qlr.font.size = Pt(11)
        qr = qd_table.rows[0].cells[1].paragraphs[0]
        qr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _date = order.date_order.date() if order.date_order else fields.Date.today()
        _date_str = _date.strftime('%d-%m-%Y')
        qrr = qr.add_run('Date: %s' % _date_str)
        qrr.font.size = Pt(11)
        qrr.bold = True
        doc.add_paragraph('')

        # To block
        to_para = doc.add_paragraph()
        to_para.add_run('To,').bold = True

        # Contact person first (if available)
        if self.contact_person:
            cp_para = doc.add_paragraph()
            cp_para.add_run(self.contact_person)

        # Company name bold - strip HTML
        import re as _re
        clean_buyer = _re.sub(r'<[^>]+>', '', self.buyer_name or '').strip()
        company_para = doc.add_paragraph()
        company_run = company_para.add_run(clean_buyer)
        company_run.bold = True
        company_run.font.size = Pt(11)



        # Full address with GST
        p = order.partner_id
        # Street
        if p.street:
            doc.add_paragraph(p.street + (', ' + p.street2 if p.street2 else ''))
        elif p.street2:
            doc.add_paragraph(p.street2)
        # City, Zip, State, Country in one line
        loc_parts = [x for x in [
            ('%s %s' % (p.city or '', p.zip or '')).strip(),
            p.state_id.name if p.state_id else '',
            p.country_id.name if p.country_id else '',
        ] if x]
        if loc_parts:
            doc.add_paragraph(', '.join(loc_parts))
        # GST bold


        # Blank line after GST
        doc.add_paragraph('')
        # Email and Phone
        if order.partner_id.email:
            doc.add_paragraph(order.partner_id.email)
        if order.partner_id.phone:
            doc.add_paragraph(order.partner_id.phone)
        # Subject
        subj_para = doc.add_paragraph()
        subj_para.add_run('Subject: ').bold = True
        subj_para.add_run(self.subject or 'Quotation for Products / Services')
        doc.add_paragraph('')
        doc.add_paragraph('Dear Sir,')
        best_offer = self.best_offer_for or ''
        doc.add_paragraph('With reference to your discussion with the undersigned as regards your subject requirement, we are pleased to quote our best offer for %s.' % best_offer)
        doc.add_paragraph('')

        # Images - Page 3
        if self.quote_image_ids:
            # Product Images heading removed
            imgs_list = [att for att in self.quote_image_ids if att.datas]
            # 2-column grid using table
            for i in range(0, len(imgs_list), 2):
                row_atts = imgs_list[i:i+2]
                img_table = doc.add_table(rows=1, cols=2)
                for j, att in enumerate(row_atts):
                    try:
                        img_buf = io.BytesIO(base64.b64decode(att.datas))
                        cell_para = img_table.rows[0].cells[j].paragraphs[0]
                        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell_para.add_run().add_picture(img_buf, width=Inches(3))
                    except Exception:
                        pass
                doc.add_paragraph('')

        # Technical Specs - New Page
        if self.technical_specs_html:
            doc.add_page_break()
            ts_heading = doc.add_paragraph()
            ts_run = ts_heading.add_run('Technical Specifications')
            ts_run.bold = True
            ts_run.font.size = Pt(11)
            # Parse HTML and render BOTH text and tables in order
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(str(self.technical_specs_html), 'html.parser')
                # Process all top-level elements in order
                for element in soup.children:
                    if hasattr(element, 'name') and element.name == 'table':
                        # Render table
                        rows = element.find_all('tr')
                        if not rows:
                            continue
                        max_cols = max(len(r.find_all(['td','th'])) for r in rows)
                        if max_cols == 0:
                            continue
                        docx_table = doc.add_table(rows=0, cols=max_cols)
                        docx_table.style = 'Table Grid'
                        for row in rows:
                            cells = row.find_all(['td', 'th'])
                            row_cells = docx_table.add_row().cells
                            for i, cell in enumerate(cells):
                                if i < max_cols:
                                    row_cells[i].text = cell.get_text(strip=True)
                                    if cell.name == 'th':
                                        for para in row_cells[i].paragraphs:
                                            for run in para.runs:
                                                run.bold = True
                                        tc = row_cells[i]._tc
                                        tcPr = tc.get_or_add_tcPr()
                                        shd = OxmlElement('w:shd')
                                        shd.set(qn('w:fill'), 'F0F0F0')
                                        shd.set(qn('w:color'), 'auto')
                                        shd.set(qn('w:val'), 'clear')
                                        tcPr.append(shd)
                        doc.add_paragraph('')
                    elif hasattr(element, 'name') and element.name in ['p', 'div']:
                        # Check if contains nested table
                        nested_tables = element.find_all('table')
                        if nested_tables:
                            # Render any text before table
                            pre_text = element.get_text(separator='\n').strip()
                            for nested_table in nested_tables:
                                pre_text = pre_text.replace(nested_table.get_text(separator='\n'), '').strip()
                            if pre_text:
                                doc.add_paragraph(pre_text)
                            # Render nested tables
                            for nested_table in nested_tables:
                                rows = nested_table.find_all('tr')
                                if not rows:
                                    continue
                                max_cols = max(len(r.find_all(['td','th'])) for r in rows)
                                if max_cols == 0:
                                    continue
                                docx_table = doc.add_table(rows=0, cols=max_cols)
                                docx_table.style = 'Table Grid'
                                for row in rows:
                                    cells = row.find_all(['td', 'th'])
                                    row_cells = docx_table.add_row().cells
                                    for i, cell in enumerate(cells):
                                        if i < max_cols:
                                            row_cells[i].text = cell.get_text(strip=True)
                                            for para in row_cells[i].paragraphs:
                                                for run in para.runs:
                                                    run.font.size = Pt(11)
                                            if cell.name == 'th':
                                                for para in row_cells[i].paragraphs:
                                                    for run in para.runs:
                                                        run.bold = True
                                                        run.font.size = Pt(11)
                                                tc = row_cells[i]._tc
                                                tcPr = tc.get_or_add_tcPr()
                                                shd = OxmlElement('w:shd')
                                                shd.set(qn('w:fill'), 'F0F0F0')
                                                shd.set(qn('w:color'), 'auto')
                                                shd.set(qn('w:val'), 'clear')
                                                tcPr.append(shd)
                                doc.add_paragraph('')
                        else:
                            # Plain text paragraph
                            text = element.get_text(strip=True)
                            if text:
                                para = doc.add_paragraph(text)
                                if element.name in ['h1','h2','h3','h4']:
                                    para.runs[0].bold = True
                    elif hasattr(element, 'name') and element.name in ['h1','h2','h3','h4']:
                        text = element.get_text(strip=True)
                        if text:
                            p = doc.add_paragraph(text)
                            p.runs[0].bold = True
                    elif hasattr(element, 'name') and element.name and element.get_text(strip=True):
                        text = element.get_text(strip=True)
                        if text:
                            doc.add_paragraph(text)
            except Exception as e:
                import logging
                logging.getLogger(__name__).error('Tech specs render error: %s', e)
                doc.add_paragraph(html2plaintext(self.technical_specs_html))


        # Table - on its own page only if tech specs or images exist
        if self.technical_specs_html or self.quote_image_ids:
            doc.add_page_break()
        quot_heading = doc.add_heading('Quotation', 2)
        quot_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Build note map for DOCX
        docx_note_map = {}
        prev_line_id = None
        for l in order.order_line:
            if l.display_type in ('line_note', 'note'):
                if prev_line_id:
                    docx_note_map.setdefault(prev_line_id, []).append(l.name or '')
            else:
                prev_line_id = l.id
        order_lines = order.order_line.filtered(lambda l: not l.display_type)
        has_discount = any(l.discount for l in order_lines)
        has_overall_discount = getattr(order, 'x_flat_discount_pct', 0) or 0

        if has_discount:
            headers = ['SR No.', 'Item Description', 'Part No', 'HSN', 'Qty', 'Unit Price', 'Discount %', 'Amount']
        else:
            headers = ['SR No.', 'Item Description', 'Part No', 'HSN', 'Qty', 'Unit Price', 'Amount']

        from docx.shared import Inches as _Inches
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        # Set column widths
        if has_discount:
            col_widths = [0.4, 2.5, 0.8, 0.8, 0.4, 0.9, 0.7, 0.8]
        else:
            col_widths = [0.4, 3.0, 0.8, 0.8, 0.4, 0.9, 0.8]
        for i, width in enumerate(col_widths):
            if i < len(table.columns):
                for cell in table.columns[i].cells:
                    cell.width = _Inches(width)
        hdr_cells = table.rows[0].cells
        # Set cell margins for padding
        from docx.oxml import OxmlElement as _OEM
        from docx.oxml.ns import qn as _qnm
        def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = _OEM('w:tcMar')
            for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
                m = _OEM('w:%s' % side)
                m.set(_qnm('w:w'), str(val))
                m.set(_qnm('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)

        for i, h in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'D9E1F2')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)

        for idx, line in enumerate(order_lines, 1):
            row_cells = table.add_row().cells
            desc = line.x_product_name or line.product_id.product_tmpl_id.with_context(lang='en_US').name or ''
            if line.x_make:
                desc += '\nMake: ' + line.x_make
            if hasattr(line, 'x_notes') and line.x_notes:
                clean_note = (line.x_notes or '').replace('&nbsp;', ' ').strip()
                if clean_note:
                    desc += '\nDescription: ' + clean_note
            for n in docx_note_map.get(line.id, []):
                desc += '\n[NOTE]' + n + '[/NOTE]'
            part_no = line.x_product_code or line.product_id.default_code or ''
            hsn = line.product_id.l10n_in_hsn_code or ''
            unit_price = line.price_unit or 0
            disc_pct = line.discount or 0
            qty = int(line.product_uom_qty or 0)
            amount = line.price_subtotal or 0

            if has_discount:
                disc_str = '(%s%%)' % int(disc_pct) if disc_pct else '-'
                row_data = [str(idx), desc, part_no, hsn, str(qty), _indian_format(unit_price), disc_str, _indian_format(amount)]
            else:
                row_data = [str(idx), desc, part_no, hsn, str(qty), _indian_format(unit_price), _indian_format(amount)]

            for i, val in enumerate(row_data):
                cell = row_cells[i]
                set_cell_margins(cell)
                if i == 1:
                    # Item description with bold labels
                    cell.paragraphs[0].clear()
                    parts = val.split('\n')
                    for pi, part in enumerate(parts):
                        if pi == 0:
                            p = cell.paragraphs[0]
                        else:
                            p = cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        if part.startswith('[NOTE]') and '[/NOTE]' in part:
                            note_text = part.replace('[NOTE]', '').replace('[/NOTE]', '')
                            r1 = p.add_run('Description:')
                            r1.bold = True
                            r1.font.size = Pt(11)
                            r2 = p.add_run(' ' + note_text)
                            r2.italic = True
                            r2.font.size = Pt(11)
                        elif part.startswith('Make:') or part.startswith('Description:'):
                            label, _, rest = part.partition(':')
                            r1 = p.add_run(label + ':')
                            r1.bold = True
                            r1.font.size = Pt(11)
                            r2 = p.add_run(rest)
                            r2.font.size = Pt(11)
                        else:
                            r = p.add_run(part)
                            r.font.size = Pt(11)
                else:
                    cell.text = val
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Totals below table - simple paragraph format like PDF
        untaxed = order.amount_untaxed
        overall_disc_pct_d = getattr(order, 'x_flat_discount_pct', 0) or 0
        overall_disc_amt_d = untaxed * overall_disc_pct_d / 100 if overall_disc_pct_d else 0
        net_d = untaxed - overall_disc_amt_d
        tax_rates_d = set()
        for line in order.order_line.filtered(lambda x: not x.display_type):
            for tax in line.tax_ids:
                tax_rates_d.add(float(tax.amount))
        total_tax_rate_d = sum(tax_rates_d)
        gst_on_d = self.x_gst_included
        grand_total_d = net_d + (net_d * total_tax_rate_d / 100) if (gst_on_d and tax_rates_d) else net_d

        def _add_total_para(doc, label, value, bold=False):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            text = '%s %s' % (label, value) if value else label
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            r.bold = bold
        _add_total_para(doc, 'Total Amount:', _indian_format(untaxed))
        if overall_disc_pct_d:
            _add_total_para(doc, 'Overall Discount (%s%%):' % int(overall_disc_pct_d), _indian_format(overall_disc_amt_d))
        if gst_on_d and tax_rates_d:
            _add_total_para(doc, 'GST (%s%%):' % int(total_tax_rate_d), '')
        final_label_d = 'Final Discount Amount:' if overall_disc_pct_d else 'Final Total Amount:'
        _add_total_para(doc, final_label_d, _indian_format(grand_total_d), bold=True)

        # Terms & Conditions - new page if no tech specs or images
        if self.selected_term_ids or order.note:
            if not (bool(re.sub(r'<[^>]+>', '', str(self.technical_specs_html or '')).strip()) or bool(self.quote_image_ids)):
                doc.add_page_break()
            terms_heading = doc.add_heading('Terms & Conditions', 2)
            terms_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if self.selected_term_ids:
                order_note = '<ol>' + ''.join('<li>%s</li>' % (t.content or '') for t in self.selected_term_ids.sorted('sequence')) + '</ol>'
            else:
                order_note = str(order.note)
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(order_note, 'html.parser')
                lists = soup.find_all(['ol', 'ul'])
                if lists:
                    for element in soup.children:
                        if hasattr(element, 'name') and element.name in ['ol', 'ul']:
                            for li in element.find_all('li', recursive=False):
                                text = li.get_text(strip=True)
                                if text:
                                    doc.add_paragraph(text, style='List Number' if element.name == 'ol' else 'List Bullet')
                        elif hasattr(element, 'name') and element.name in ['p', 'div']:
                            text = element.get_text(strip=True)
                            if text:
                                doc.add_paragraph(text)
                else:
                    for element in soup.find_all(['p', 'div']):
                        text = element.get_text(strip=True)
                        if text:
                            doc.add_paragraph(text)
                    if not soup.find_all(['p', 'div']):
                        doc.add_paragraph(soup.get_text(strip=True))
            except Exception as e:
                import logging
                logging.getLogger(__name__).error('Terms render error: %s', e)
                doc.add_paragraph(html2plaintext(order.note))



        # ── Closing signature block ──
        doc.add_paragraph('')
        doc.add_paragraph('Thanking You,')
        doc.add_paragraph('Sincerely,')
        doc.add_paragraph('')
        sig_para = doc.add_paragraph()
        sig_run = sig_para.add_run((order.user_id.name or '') if order.user_id else '')
        sig_run.bold = True
        comp_para = doc.add_paragraph()
        comp_run = comp_para.add_run(order.company_id.name or '')
        comp_run.bold = True
        # Save
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        docx_data = buf.read()

        attachment = self.env['ir.attachment'].create({
            'name': 'Quotation-%s.docx' % (self.quote_name or 'quote'),
            'type': 'binary',
            'datas': base64.b64encode(docx_data).decode(),
            'res_model': 'sale.quote.preview.wizard',
            'res_id': self.id,
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        })
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/content/%s?download=true' % attachment.id,
            'target': 'new',
        }

    def action_send_email(self):
        self.ensure_one()
        # Save selected terms to order note
        if self.selected_term_ids:
            items = ''.join('<li>%s</li>' % (t.content or '') for t in self.selected_term_ids.sorted('sequence'))
            self.order_id.sudo().write({'note': '<ol>%s</ol>' % items})

        import base64
        self._rebuild_document_html()
        order = self.order_id
        best_offer = self.best_offer_for or ''
        subject = self.subject or ('Quotation %s - %s' % (self.quote_name or '', order.partner_id.name or ''))
        body = '<p>Dear Sir,</p><p>' + self.env['ir.config_parameter'].sudo().get_param('sale.quote.intro.template', 'With reference to your discussion with the undersigned as regards your subject requirement, we are pleased to quote our best offer for') + '</p><p>Please find the attached quotation for your reference.</p><p>Best Regards,<br/>%s</p>' % (self.seller_name or '')

        # Generate PDF from HTML content directly
        try:
            from odoo.tools import pdf as pdf_tools
            html = self.document_html or ''
            # Use wkhtmltopdf directly
            import subprocess, tempfile, os
            with tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8') as f:
                f.write('<html><body>%s</body></html>' % str(html))
                html_file = f.name
            pdf_file = html_file.replace('.html', '.pdf')
            subprocess.run(['wkhtmltopdf', '--quiet', '--no-header-line', '--no-footer-line', '--margin-top', '10', '--margin-bottom', '10', '--margin-left', '10', '--margin-right', '10', html_file, pdf_file], check=True)
            with open(pdf_file, 'rb') as f:
                pdf_content = f.read()
            os.unlink(html_file)
            os.unlink(pdf_file)
            attachment = self.env['ir.attachment'].create({
                'name': 'Quotation-%s.pdf' % (self.quote_name or 'quote'),
                'type': 'binary',
                'datas': base64.b64encode(pdf_content).decode(),
                'res_model': 'sale.order',
                'res_id': order.id,
                'mimetype': 'application/pdf',
            })
            attachment_ids = [(4, attachment.id)]
        except Exception as e:
            import logging
            logging.getLogger(__name__).error('PDF generation failed: %s', e)
            attachment_ids = []

        return {
            'type': 'ir.actions.act_window',
            'name': 'Send Email',
            'res_model': 'mail.compose.message',
            'view_mode': 'form',
            'target': 'new',
            'context': {
                'default_model': 'sale.order',
                'default_res_ids': [order.id],
                'default_subject': subject,
                'default_body': body,
                'default_attachment_ids': attachment_ids,
                'default_partner_ids': [(4, order.partner_id.id)] if order.partner_id else [],
            },
        }

    def action_print_edited_pdf(self):
        self.ensure_one()
        import logging
        _log = logging.getLogger(__name__)
        # Save selected terms to order note
        if self.selected_term_ids:
            items = ''.join('<li>%s</li>' % (t.content or '') for t in self.selected_term_ids.sorted('sequence'))
            self.order_id.sudo().write({'note': '<ol>%s</ol>' % items})
        # Always rebuild to include technical specs and images
        self._rebuild_document_html()
        _log.warning("PDF HTML length: %s", len(str(self.document_html or '')))
        _log.warning("Has Quotation title: %s", '>Quotation<' in str(self.document_html or ''))
        _log.warning("Has page-break: %s", 'page-break-before:always' in str(self.document_html or ''))
        return self.env['ir.actions.report'].search([('report_name','=','custom_crm_extended.report_sale_quote_preview_wizard')], limit=1).report_action(self)

